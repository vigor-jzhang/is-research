"""Deterministic manuscript exporters: Markdown, LaTeX, DOCX, PDF.

All exporters render from the immutable FormattedManuscript and produce
bytes; content is stored via the BlobStore. Rendering is deterministic
(no timestamps, no randomness).
"""

from __future__ import annotations

import html
import re
from typing import Any

from research_harness.contracts.manuscript_exporter import (
    ExportPayload,
    ExportPayloadBuilder,
    ManuscriptExporter,
)
from research_harness.research.schemas.publication import FormattedManuscript

_MD_LATEX_SPECIALS = {
    "\\": r"\textbackslash{}",
    "{": r"\{",
    "}": r"\}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "^": r"\textasciicircum{}",
    "~": r"\textasciitilde{}",
}


def latex_escape(text: str) -> str:
    out: list[str] = []
    for ch in text:
        out.append(_MD_LATEX_SPECIALS.get(ch, ch))
    return "".join(out)


class MarkdownExporter:
    format = "markdown"
    renderer = "markdown"
    renderer_version = "1.0"

    def render(self, manuscript: FormattedManuscript, profile: Any) -> ExportPayload:
        lines: list[str] = []
        fm = manuscript.front_matter
        lines.append(f"# {fm.title}")
        lines.append("")
        if fm.abstract:
            lines.append("## Abstract")
            lines.append("")
            lines.append(fm.abstract)
            lines.append("")
        if fm.keywords:
            lines.append(f"**Keywords:** {', '.join(fm.keywords)}")
            lines.append("")
        for section in manuscript.sections:
            lines.append(f"## {section.title}")
            lines.append("")
            lines.append(section.body)
            lines.append("")
        if manuscript.bibliography is not None and manuscript.bibliography.rendered_text:
            lines.append("## References")
            lines.append("")
            lines.append(manuscript.bibliography.rendered_text)
        data = ("\n".join(lines) + "\n").encode()
        return ExportPayloadBuilder.build(
            self.format, self.renderer, self.renderer_version, data, "text/markdown"
        )


class LatexExporter:
    format = "latex"
    renderer = "latex"
    renderer_version = "1.0"

    def render(self, manuscript: FormattedManuscript, profile: Any) -> ExportPayload:
        docclass = "article"
        if profile is not None:
            rules = getattr(profile, "formatting_rules", {}) or {}
            docclass = rules.get("latex_documentclass", docclass)
        lines = [
            "\\documentclass{" + docclass + "}",
            "\\usepackage[utf8]{inputenc}",
            "\\title{" + latex_escape(manuscript.front_matter.title) + "}",
            "\\date{}",
        ]
        if manuscript.anonymous_review:
            lines.append("\\author{}")
        else:
            authors = ", ".join(manuscript.front_matter.authors) or "Anonymous"
            lines.append("\\author{" + latex_escape(authors) + "}")
        lines.append("\\begin{document}")
        lines.append("\\maketitle")
        if manuscript.front_matter.abstract:
            lines.append("\\begin{abstract}")
            lines.append(latex_escape(manuscript.front_matter.abstract))
            lines.append("\\end{abstract}")
        for section in manuscript.sections:
            lines.append("\\section{" + latex_escape(section.title) + "}")
            lines.append(latex_escape(section.body))
        if manuscript.bibliography is not None and manuscript.bibliography.entries:
            lines.append("\\begin{thebibliography}{99}")
            for entry in manuscript.bibliography.entries:
                lines.append(
                    f"\\bibitem{{{entry.paper_identity_id}}} {latex_escape(entry.rendered)}"
                )
            lines.append("\\end{thebibliography}")
        lines.append("\\end{document}")
        data = ("\n".join(lines) + "\n").encode()
        return ExportPayloadBuilder.build(
            self.format, self.renderer, self.renderer_version, data, "application/x-tex"
        )


class DocxExporter:
    format = "docx"
    renderer = "python-docx"
    renderer_version = "1.2"

    def render(self, manuscript: FormattedManuscript, profile: Any) -> ExportPayload:
        import docx  # type: ignore[import-untyped]

        document = docx.Document()
        document.add_heading(manuscript.front_matter.title, level=0)
        if manuscript.front_matter.abstract:
            document.add_heading("Abstract", level=1)
            document.add_paragraph(manuscript.front_matter.abstract)
        if manuscript.front_matter.keywords:
            document.add_paragraph("Keywords: " + ", ".join(manuscript.front_matter.keywords))
        for section in manuscript.sections:
            document.add_heading(section.title, level=1)
            document.add_paragraph(section.body)
        if manuscript.bibliography is not None and manuscript.bibliography.rendered_text:
            document.add_heading("References", level=1)
            document.add_paragraph(manuscript.bibliography.rendered_text)
        import io

        buf = io.BytesIO()
        document.save(buf)
        return ExportPayloadBuilder.build(
            self.format,
            self.renderer,
            self.renderer_version,
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


class PdfExporter:
    format = "pdf"
    renderer = "reportlab"
    renderer_version = "5.0"

    def render(self, manuscript: FormattedManuscript, profile: Any) -> ExportPayload:
        import io

        from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
        from reportlab.lib.styles import (  # type: ignore[import-untyped]
            ParagraphStyle,
            getSampleStyleSheet,
        )
        from reportlab.lib.units import cm  # type: ignore[import-untyped]
        from reportlab.platypus import (  # type: ignore[import-untyped]
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title=manuscript.front_matter.title,
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "body", parent=styles["Normal"], leading=14, spaceAfter=6, alignment=4
        )
        story: list[Any] = []
        story.append(Paragraph(html.escape(manuscript.front_matter.title), styles["Title"]))
        if manuscript.front_matter.abstract:
            story.append(Spacer(1, 8))
            story.append(Paragraph("Abstract", styles["Heading2"]))
            story.append(Paragraph(html.escape(manuscript.front_matter.abstract), body))
        for section in manuscript.sections:
            story.append(Spacer(1, 10))
            story.append(Paragraph(html.escape(section.title), styles["Heading2"]))
            for para in re.split(r"\n\s*\n", section.body.strip()):
                story.append(Paragraph(html.escape(para), body))
        if manuscript.bibliography is not None and manuscript.bibliography.rendered_text:
            story.append(Spacer(1, 10))
            story.append(Paragraph("References", styles["Heading2"]))
            story.append(Paragraph(html.escape(manuscript.bibliography.rendered_text), body))
        doc.build(story)
        return ExportPayloadBuilder.build(
            self.format, self.renderer, self.renderer_version, buf.getvalue(), "application/pdf"
        )


EXPORTERS: list[ManuscriptExporter] = [
    MarkdownExporter(),
    LatexExporter(),
    DocxExporter(),
    PdfExporter(),
]


def get_exporter(fmt: str) -> ManuscriptExporter:
    for exporter in EXPORTERS:
        if exporter.format == fmt:
            return exporter
    raise ValueError(
        f"unknown export format {fmt!r} (available: {', '.join(e.format for e in EXPORTERS)})"
    )


def available_formats() -> list[str]:
    return [e.format for e in EXPORTERS]
